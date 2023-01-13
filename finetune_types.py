from PIL import Image
import docx2txt
import os, os.path
from docx import Document 
from  builtins import any as b_any
import matplotlib.pyplot as plt
from sklearn.metrics import confusion_matrix, classification_report
import glob
from torchsummary import summary
from tqdm.notebook import tqdm
import pandas as pd
import albumentations as A
from sklearn.model_selection import train_test_split
import torch
import torchvision.transforms as transforms
import torchvision
import time
from albumentations.pytorch import ToTensorV2
from sklearn.model_selection import KFold
from torch.utils.data import Dataset, DataLoader,WeightedRandomSampler
import numpy as np
import copy
import cv2
import torch.nn as nn
from torch.utils.data.sampler import SubsetRandomSampler
import torch.nn.functional as F
from torch.nn import Module
from torch.nn import Conv2d
from torch.nn import Linear
from torch.nn import MaxPool2d
from torch.nn import ReLU
from torch.nn import LogSoftmax
from torch import flatten
import gc
from collections import Counter
from sklearn.metrics import top_k_accuracy_score

device = 'cuda' if torch.cuda.is_available() else 'cpu'
def report_gpu():
   print(torch.cuda.list_gpu_processes())
   gc.collect()
   torch.cuda.empty_cache()

document = Document("veri5_fused/high_low_level_labels.docx")
tables = document.tables
bold_list = []
for table in tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.bold and run.text!=" ":
                       
                        bold_list.append(run.text)
                        
                        

print("Number of pollen families: ",len(bold_list))
'''
for pollen in bold_list:
    print(pollen)
'''
pollens = {}
with open(r"veri5_fused/high_low_level_labels.txt", 'r',encoding="utf8") as fp:
    lines = []
    pollen_types=[]
    all = []
  
    for i, line in enumerate(fp):
        if line.strip()[1:] in bold_list:
            pollen_fam = line.strip()[1:]
            lines.append(line.strip()[1:])
            pollen_types=[]
            continue
            
        elif line.strip() in bold_list:
            pollen_fam = line.strip()
            lines.append(line.strip())
            pollen_types=[]
            continue
        else:
            #pol_typ_dict = {}
            #pol_typ_dict[line.strip()] = 0
            pollen_types.append(line.strip())
        all.append(line.strip())

        pollens[pollen_fam] = pollen_types
        
        

print(len(lines))
print(lines==bold_list)
print(pollens["Acanthaceae"])
print(len(all))



rootdir = 'veri5_fused/data'
pollens_imgs={}
for subdir, dirs, files in os.walk(rootdir):
    pol_type = subdir[17:]
    #print("Pollen type: "+pol_type)
    count = 0
    for file in files:
        count+=1
    pollens_imgs[pol_type] = count
    if "" in pollens_imgs and pol_type=="":
        del pollens_imgs[""]
        continue
pollens_imgs2={}
for subdir, dirs, files in os.walk(rootdir):
    pol_type = subdir[17:]
    #print("Pollen type: "+pol_type)
    imgs=[]
    for file in files:
        
        imgs.append(subdir+"/"+file)
    pollens_imgs2[pol_type] = imgs
    if "" in pollens_imgs2 and pol_type=="":
        del pollens_imgs2[""]
        continue


print(pollens_imgs["Xeranthemum longipopposum"])
print(pollens_imgs2["Xeranthemum longipopposum"])

lists = sorted(pollens_imgs.items()) # sorted by key, return a list of tuples
x, y = zip(*lists) # unpack a list of pairs into two tuples
print(x,y)
plt.plot(x, y)
plt.show()


base = "/home/deren/Desktop/veri5_fused/data/"
folders = glob.glob(base + "*")
folders = [x.replace(base,"") for x in folders]
#print("folders:", folders)

def str_comp(str1, str2):

    match = 0
    substr = 0
    for letter in str1:
        if letter in str2:
            match += 1

    for item1 in str1.split(" "):
        for item2 in str2.split(" "):
            if(item1 == item2 and len(item1) == len(item2)):
                #if(str1.find("Eremurus") != -1):
                    #print("str1:", str1, "str2:", str2, "item1:", item1, "item2:", item2, item1==item2)
                substr += 0.1

    if(str1[:-1] == str2):
        substr += 0.1

    return (match / len(str1)) + substr

for fam in pollens.keys():
    for gen in pollens[fam]:
        my_sims = {}
        for folder in folders:
            folder = folder.replace(base,"")
            my_sims[folder] = str_comp(gen,folder)
        sorted_sims = sorted(my_sims.items(), key=lambda x:x[1], reverse=True)
        #print("gen:", gen, "my_sims:", my_sims)
        #print("gen:", gen, "folder:", folder, "similarity:", str_comp(gen, folder))
        print("gen:", gen, "Sorted:", sorted_sims[:5])


def class_compar(type, fam):

    scores = []
    for item in fam:
        scores.append(str_comp(type, item))

    return max(scores)

classes = {}

for folder in folders:
    my_classes = {}
    for fam in pollens.keys():
        my_classes[fam] = class_compar(folder, pollens[fam])
        sorted_classes = sorted(my_classes.items(), key=lambda x:x[1], reverse=True)
    print("type:", folder, "sorted_classes:", sorted_classes[0][0])
    try:
        classes[sorted_classes[0][0]]["types"].append(folder)
    except:
        classes[sorted_classes[0][0]] = {}
        classes[sorted_classes[0][0]]["types"] = []
        classes[sorted_classes[0][0]]["types"].append(folder)



for key in classes:
    classes[key]["tot_img"] = 0
    for _type in classes[key]["types"]:
        classes[key]["tot_img"] += pollens_imgs[_type]


print("##########################################")
print("classes:", classes)

tot_imgs = {}
for key in classes:
    tot_imgs[key] = classes[key]["tot_img"]

print("##########################################")
print("tot_imgs:", tot_imgs)

new = pd.DataFrame(tot_imgs, index=[0])
print(new)
data = []
for key in classes:
    
    for _type in classes[key]["types"]:
        for img in pollens_imgs2[_type]:
            dct = {}
            
            
            #dct["img_name"] = torchvision.io.read_image(img)
            dct["image_paths"]=img
            dct["pol_type"] = _type
            dct["pol_fam"] = key
            data.append(dct)

img_paths = []
pol_t=[]
for key in classes:
    
    for _type in classes[key]["types"]:
        for img in pollens_imgs2[_type]:
            
            #dct["img_name"] = torchvision.io.read_image(img)
            img_paths.append(img)
            pol_t.append(_type)


idx_to_class = {i:j for i, j in enumerate(pollens_imgs.keys())}
class_to_idx = {value:key for key,value in idx_to_class.items()}


X_train, X_test, y_train, y_test = train_test_split(img_paths, pol_t,stratify=pol_t, test_size=0.2, random_state=42)
print(X_train[0:5],y_train[0:5])
print(X_test[0:5],y_test[0:5])

X_train, X_valid, y_train, y_valid = train_test_split(X_train, y_train,stratify=y_train, test_size=0.2, random_state=42)
print(X_train[0:5],y_train[0:5])
print(X_test[0:5],y_test[0:5])
train_transforms = A.Compose(
        [
            A.Resize(height=256, width=256),
            A.SmallestMaxSize(max_size=350),
            A.ShiftScaleRotate(shift_limit=0.05, scale_limit=0.05, rotate_limit=360, p=0.5),
            A.RandomCrop(height=224, width=224),
            A.Normalize(mean=(0.485, 0.456, 0.406), std=(0.229, 0.224, 0.225)),
            ToTensorV2(),
        ]
    )

test_transforms = A.Compose(
    [
        A.Resize(height=256, width=256),
        A.CenterCrop(height=224, width=224),
        A.Normalize(mean=(0.485, 0.456, 0.406), std=(0.229, 0.224, 0.225)),
        ToTensorV2(),
    ]
)
class PollenTypeDataset(Dataset):
    def __init__(self, image_paths, pol_t,transform=False):
        self.image_paths = image_paths
        self.transform = transform
        self.labels = pol_t
        
    def __len__(self):
        return len(self.image_paths)

    def __getitem__(self, idx):
        image_filepath = self.image_paths[idx]
        image = cv2.imread(image_filepath)
        image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        
        label = self.labels[idx]
        label = class_to_idx[label]
        if self.transform is not None:
            image = self.transform(image=image)["image"]
        
        return image, label
    
#######################################################
#                  Create Dataset
#######################################################

#train_dataset = PollenTypeDataset(img_paths,pol_t,train_transforms)
train_dataset = PollenTypeDataset(X_train,y_train,test_transforms)
valid_dataset = PollenTypeDataset(X_valid,y_valid,test_transforms)
test_dataset = PollenTypeDataset(X_test,y_test,test_transforms)

print('The shape of tensor for 50th image in train dataset: ',train_dataset[49][0].shape)
print('The shape of tensor for 50th image in train dataset: ',valid_dataset[49][0].shape)
print('The shape of tensor for 50th image in train dataset: ',test_dataset[49][0].shape)

target_list = []
for _, t in train_dataset:
    target_list.append(t)
    
target_list = torch.tensor(target_list)
class_dist = Counter(y_train).values()
print(Counter(y_train).keys()) 
print(class_dist) 
class_count = [i for i in class_dist]
print(class_count)
class_weights = 1./torch.tensor(class_count, dtype=torch.float) 
print(class_weights)
class_weights_all = class_weights[target_list]
weighted_sampler = WeightedRandomSampler(
    weights=class_weights_all,
    num_samples=len(class_weights_all),
    replacement=True
)
EPOCHS = 100
BATCH_SIZE = 12
LEARNING_RATE = 0.0007


train_loader = DataLoader(dataset=train_dataset,
                          batch_size=BATCH_SIZE,
                          sampler=weighted_sampler
)
val_loader = DataLoader(dataset=valid_dataset, batch_size=1)
test_loader = DataLoader(dataset=test_dataset, batch_size=1)


model = torchvision.models.resnext101_32x8d(pretrained=True)
for param in list(model.parameters())[:-10]:
    param.requires_grad = False
num_ftrs = model.fc.in_features

#model.fc = nn.Linear(num_ftrs, len(class_to_idx.items()))
model.fc = nn.Sequential(
    nn.Dropout(0.5),
    nn.Linear(num_ftrs, len(class_to_idx.items()))
)

model.to(device)


criterion = nn.CrossEntropyLoss(weight=class_weights.to(device))
optimizer = torch.optim.Adam(model.parameters(), lr=LEARNING_RATE)

def multi_acc(y_pred, y_test):
    y_pred_softmax = torch.log_softmax(y_pred, dim = 1)
    
    _, y_pred_tags = torch.max(y_pred_softmax, dim = 1)    
    
    correct_pred = (y_pred_tags == y_test).float()

    acc = correct_pred.sum() / len(correct_pred)
    
    acc = torch.round(acc * 100)
    
    return acc
def multi_acc_topk(y_pred, y_test,k):
    y_pred_softmax = torch.log_softmax(y_pred, dim = 1)
    #correct_pred = torch.empty(BATCH_SIZE)
   
    _, y_pred_tags = torch.topk(y_pred_softmax, dim = 1,k=k)
    #print(y_pred_tags.shape, y_test.shape)
    #print((y_pred_tags == y_test[..., None]).shape)
    #print((y_pred_tags == y_test[..., None]).any(dim=-1).shape)
    
    correct_pred = (y_pred_tags == y_test[..., None]).any(dim=-1).float()
    acc = correct_pred.sum() / len(correct_pred)
    
    #acc = correct_pred.mean()
    #exit()

    #print((y_pred_tags == y_test).any(dim=-1).float())
    #print((y_test in y_pred_tags))
    # for i in range(len(y_test)):
    #     correct_pred[i] = (y_pred_tags[i] == y_test[i]).any(dim=-1).float()
    # acc = correct_pred.sum() / len(correct_pred)
    
    acc = torch.round(acc * 100)
    #print("GGGG", acc)
    
    return acc
accuracy_stats = {
    'train': [],
    "val": []
}

accuracy_stats_3 = {
    'train': [],
    "val": []
}

accuracy_stats_5 = {
    'train': [],
    "val": []
}
loss_stats = {
    'train': [],
    "val": []
}
print("Begin training.")
for e in range(1, EPOCHS+1):
    
    # TRAINING
    train_epoch_loss = 0
    train_epoch_acc = 0
    train_epoch_acc_3 = 0
    train_epoch_acc_5 = 0
    model.train()
    for X_train_batch, y_train_batch in train_loader:
        X_train_batch, y_train_batch = X_train_batch.to(device), y_train_batch.to(device)
        optimizer.zero_grad()
        
        y_train_pred = model(X_train_batch)
        
        train_loss = criterion(y_train_pred, y_train_batch)
        train_acc = multi_acc(y_train_pred, y_train_batch)
        train_acc_3 = multi_acc_topk(y_train_pred,y_train_batch,3)
        train_acc_5 = multi_acc_topk(y_train_pred,y_train_batch,5)
        
        train_loss.backward()
        optimizer.step()
        
        train_epoch_loss += train_loss.item()
        train_epoch_acc += train_acc.item()
        train_epoch_acc_3 += train_acc_3.item()
        train_epoch_acc_5 += train_acc_5.item()
        
        
    # VALIDATION    
    with torch.no_grad():
        
        val_epoch_loss = 0
        val_epoch_acc = 0
        val_epoch_acc_3 = 0
        val_epoch_acc_5 = 0
        
        model.eval()
        for X_val_batch, y_val_batch in val_loader:
            X_val_batch, y_val_batch = X_val_batch.to(device), y_val_batch.to(device)
            
            y_val_pred = model(X_val_batch)
                        
            val_loss = criterion(y_val_pred, y_val_batch)
            val_acc = multi_acc(y_val_pred, y_val_batch)
            val_acc_3 = multi_acc_topk(y_val_pred, y_val_batch,3)
            val_acc_5 = multi_acc_topk(y_val_pred, y_val_batch,5)
            
            val_epoch_loss += val_loss.item()
            val_epoch_acc += val_acc.item()
            val_epoch_acc_3 += val_acc_3.item()
            val_epoch_acc_5 += val_acc_5.item()

    loss_stats['train'].append(train_epoch_loss/len(train_loader))
    loss_stats['val'].append(val_epoch_loss/len(val_loader))
    accuracy_stats['train'].append(train_epoch_acc/len(train_loader))
    accuracy_stats['val'].append(val_epoch_acc/len(val_loader))
    accuracy_stats_3['train'].append(train_epoch_acc_3/len(train_loader))
    accuracy_stats_3['val'].append(val_epoch_acc_3/len(val_loader))
    accuracy_stats_5['train'].append(train_epoch_acc_5/len(train_loader))
    accuracy_stats_5['val'].append(val_epoch_acc_5/len(val_loader))

    print(f'Epoch {e+0:03}: | Train Loss: {train_epoch_loss/len(train_loader):.5f} | Val Loss: {val_epoch_loss/len(val_loader):.5f} | Train Acc_1: {train_epoch_acc/len(train_loader):.3f}| Val Acc_1: {val_epoch_acc/len(val_loader):.3f} | Train Acc_3: {train_epoch_acc_3/len(train_loader):.3f} | Val Acc_3: {val_epoch_acc_3/len(val_loader):.3f} | Train Acc_5: {train_epoch_acc_5/len(train_loader):.3f} | Val Acc_5: {val_epoch_acc_5/len(val_loader):.3f}')

y_pred_list = []
y_true_list = []
with torch.no_grad():
    for x_batch, y_batch in val_loader:
        x_batch, y_batch = x_batch.to(device), y_batch.to(device)
        
        y_test_pred = model(x_batch)
        
        y_test_pred = torch.log_softmax(y_test_pred, dim=1)
        _, y_pred_tag = torch.max(y_test_pred, dim = 1)
    
        y_pred_list.append(y_pred_tag.cpu().numpy())
        y_true_list.append(y_batch.cpu().numpy())

y_pred_list = [a.squeeze().tolist() for a in y_pred_list]
y_true_list = [a.squeeze().tolist() for a in y_true_list]
print("For validation set: ")
print(classification_report(y_true_list, y_pred_list))
print(confusion_matrix(y_true_list, y_pred_list))


for i in range(len(y_test)):
    y_test[i]==class_to_idx[y_test[i]]
y_pred_list = []
y_true_list = []
with torch.no_grad():
    for x_batch, y_batch in test_loader:
        x_batch, y_batch = x_batch.to(device), y_batch.to(device)
        
        y_test_pred = model(x_batch)
        
        y_test_pred = torch.log_softmax(y_test_pred, dim=1)
        _, y_pred_tag = torch.max(y_test_pred, dim = 1)
    
        y_pred_list.append(y_pred_tag.cpu().numpy())
        y_true_list.append(y_batch.cpu().numpy())

y_pred_list = [a.squeeze().tolist() for a in y_pred_list]
y_true_list = [a.squeeze().tolist() for a in y_true_list]
print("For test set: ")
print(classification_report(y_true_list, y_pred_list))
print(confusion_matrix(y_true_list, y_pred_list))
y_pred_list = []
y_true_list = []
with torch.no_grad():
    for x_batch, y_batch in test_loader:
        x_batch, y_batch = x_batch.to(device), y_batch.to(device)
        
        y_test_pred = model(x_batch)
        
        y_test_pred = torch.log_softmax(y_test_pred, dim=1)
        #_, y_pred_tags = torch.topk(y_test_pred, dim = 1,k=3)
        y_pred_list.append(y_test_pred.cpu().numpy())
        y_true_list.append(y_batch.cpu().numpy())

print([a.squeeze().tolist() for a in y_pred_list])
print(len(y_pred_list))
y_pred_list = [a.squeeze().tolist() for a in y_pred_list]
y_true_list = [a.squeeze().tolist() for a in y_true_list]  
print(len(y_true_list))
print(len(y_pred_list))


print("Top 3 accuracy: " + top_k_accuracy_score(y_true_list, y_pred_list, k=3))


    
  
   
    
    