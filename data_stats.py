from PIL import Image
import docx2txt
import os, os.path
from docx import Document 
from  builtins import any as b_any
import matplotlib.pyplot as plt
from sklearn.metrics import confusion_matrix, classification_report
import glob
from torchsummary import summary
from tqdm.notebook import tqdm
import pandas as pd
import albumentations as A
from sklearn.model_selection import train_test_split
import torch
import torchvision.transforms as transforms
import torchvision
import time
from albumentations.pytorch import ToTensorV2
from sklearn.model_selection import KFold
from torch.utils.data import Dataset, DataLoader,WeightedRandomSampler
import numpy as np
import copy
import cv2
import torch.nn as nn
from torch.utils.data.sampler import SubsetRandomSampler
import torch.nn.functional as F
from torch.nn import Module
from torch.nn import Conv2d
from torch.nn import Linear
from torch.nn import MaxPool2d
from torch.nn import ReLU
from torch.nn import LogSoftmax
from torch import flatten
import gc
from collections import Counter
from sklearn.metrics import top_k_accuracy_score
import plotly.graph_objects as go


device = 'cuda' if torch.cuda.is_available() else 'cpu'
def report_gpu():
   print(torch.cuda.list_gpu_processes())
   gc.collect()
   torch.cuda.empty_cache()


base = "/home/deren/Desktop/veri5_fused/data/"
folders = glob.glob(base + "*")
folders = [x.replace(base,"") for x in folders]
#print("folders:", folders)

def print_this(t, t_name):

    print("##################" + t_name + "#####################")
    print(t)
    print("##################" + t_name + "#####################")


def ret_bold_list():
    document = Document("veri5_fused/high_low_level_labels.docx")
    tables = document.tables
    bold_list = []
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.bold and run.text!=" ":
                            bold_list.append(run.text)
                        
    #print("Number of pollen families: ",len(bold_list))
    return bold_list


def ret_pollens():
    pollens = {}
    with open(r"veri5_fused/high_low.txt", 'r',encoding="utf8") as fp:
        lines = []
        pollen_types=[]
        all = []
    
        for i, line in enumerate(fp):
            if line.strip()[1:] in bold_list:
                pollen_fam = line.strip()[1:]
                lines.append(line.strip()[1:])
                pollen_types=[]
                continue
                
            elif line.strip() in bold_list:
                pollen_fam = line.strip()
                lines.append(line.strip())
                pollen_types=[]
                continue
            else:
                pollen_types.append(line.strip())
            all.append(line.strip())

            pollens[pollen_fam] = pollen_types
    
    return pollens
        
    
def ret_pollens_imgses():
    rootdir = 'veri5_fused/data'
    pollens_imgs={}
    for subdir, dirs, files in os.walk(rootdir):
        pol_type = subdir[17:]
        count = 0
        for file in files:
            count+=1
        pollens_imgs[pol_type] = count
        if "" in pollens_imgs and pol_type=="":
            del pollens_imgs[""]
            continue
    pollens_imgs2={}
    for subdir, dirs, files in os.walk(rootdir):
        pol_type = subdir[17:]
        imgs=[]
        for file in files:
            
            imgs.append(subdir+"/"+file)
        pollens_imgs2[pol_type] = imgs
        if "" in pollens_imgs2 and pol_type=="":
            del pollens_imgs2[""]
            continue

    return pollens_imgs, pollens_imgs2


def str_comp(str1, str2):

    match = 0
    substr = 0
    for letter in str1:
        if letter in str2:
            match += 1

    for item1 in str1.split(" "):
        for item2 in str2.split(" "):
            if(item1 == item2 and len(item1) == len(item2)):
                #if(str1.find("Eremurus") != -1):
                    #print("str1:", str1, "str2:", str2, "item1:", item1, "item2:", item2, item1==item2)
                substr += 0.1

    if(str1[:-1] == str2):
        substr += 0.1

    try:
        return (match / len(str1)) + substr
    except:
        print("str1:", str1, "str2:", str2, "match:", match, "substr:", substr)


def ret_sorted_sims():
    for fam in pollens.keys():
        pollens[fam] = [x for x in pollens[fam] if x != ""]
        for gen in pollens[fam]:
            my_sims = {}
            for folder in folders:
                folder = folder.replace(base,"")
                my_sims[folder] = str_comp(gen,folder)
            sorted_sims = sorted(my_sims.items(), key=lambda x:x[1], reverse=True)
            #print("gen:", gen, "Sorted:", sorted_sims[:5])

    return sorted_sims


def class_compar(type, fam):

    scores = []
    for item in fam:
        scores.append(str_comp(type, item))
        

    return max(scores)


def ret_classes():
    classes = {}

    for folder in folders:
        my_classes = {}
        for fam in pollens.keys():
            my_classes[fam] = class_compar(folder, pollens[fam])
            sorted_classes = sorted(my_classes.items(), key=lambda x:x[1], reverse=True)
        #print("folder:", folder, "sorted classes:", sorted_classes)

        try:
            classes[sorted_classes[0][0]]["types"].append(folder)
        except:
            classes[sorted_classes[0][0]] = {}
            classes[sorted_classes[0][0]]["types"] = []
            classes[sorted_classes[0][0]]["types"].append(folder)
    
    return classes



def ret_imgpath_lists():
    img_paths = []
    pol_t=[]

    for key in classes:
        for _type in classes[key]["types"]:
            for img in pollens_imgs2[_type]:
                
                #dct["img_name"] = torchvision.io.read_image(img)
                img_paths.append(img)
                pol_t.append(_type)
    return img_paths,pol_t


def swap(classes, _type, old_owner, new_owner):

    classes[old_owner]["types"].remove(_type)
    classes[new_owner]["types"].append(_type)

    return classes

def corrections(classes):

    ##Manual arrangements due to faulties in dataset
    classes["Boraginaceae"]["types"].remove("Rosaceae")
    classes["Boraginaceae"]["types"].remove("Poaceae")
    classes['Amaranthaceae']["types"].remove("Chenopodium sp")
    classes["Poaceae"] = {"types":["Poaceae"]}
    classes["Chenopodiaceae"] = {"types":["Chenopodium sp"]}
    classes["Rosaceae"]["types"].append("Rosaceae")

    classes = swap(classes, "Circium arvense", "Campanulaceae", "Asteraceae")
    classes = swap(classes, "Apiaceae sp", "Campanulaceae", "Apiaceae")

    return classes

def treemap(classes, pollens_imgs, tot_imgs):

   tvalues = []
   tlabels = []
   tparents = []
   
   for key in list(classes.keys())[:25]:
      val = 0
      for pol in classes[key]["types"]:
         val += pollens_imgs[pol]
         print("Class:", key, "pollen:", pol, "val:", pollens_imgs[pol])
      tvalues.append(val)
      tlabels.append(key)
      tparents.append("")

   #val = 0
   #for i in range(1, len(tvalues)):
   #   val += tvalues[i]
   #tvalues[0] = val

   for key in list(classes.keys())[:25]:
      for pol in classes[key]["types"]:
         tvalues.append(pollens_imgs[pol])
         tlabels.append(pol)
         tparents.append(key)

   fig = go.Figure(go.Sunburst(
      labels = tlabels,
      values = tvalues,
      parents = tparents,
      marker_colorscale = 'Blues',
      branchvalues= 'total'
   ))
   default_layout = {
    'margin': {'r': 5, 't': 20, 'l': 5, 'b': 5}
   }
   fig.update_layout(
                    **default_layout,
                    template='simple_white',
                    )
   print("values:", tvalues)
   print("labels:", tlabels)
   print("parents:", tparents)
   fig.update_layout(margin = dict(t=15, l=5, r=5, b=5))
   fig.show()
   go.Figure.write_html(fig, "sunb.html")

def treemap2():
   fig = go.Figure(go.Treemap(
      labels = ["Eve","Cain", "Seth", "Enos", "Noam", "Abel", "Awan", "Enoch", "Azura"],
      parents = ["", "Eve", "Eve", "Seth", "Seth", "Eve", "Eve", "Awan", "Eve"],
      root_color="lightgrey"
   ))

   fig.update_layout(margin = dict(t=50, l=25, r=25, b=25), autosize=False)
   fig.show()
         
if __name__ == "__main__":

    bold_list = ret_bold_list()
    pollens = ret_pollens()
    pollens_imgs, pollens_imgs2 = ret_pollens_imgses()
  
    sorted_sims = ret_sorted_sims()
    classes = ret_classes()
    classes = corrections(classes)
    print_this(classes, "classes")
    img_paths, pol_t = ret_imgpath_lists()
    print("Number of pollen families: ",len(classes.keys()))
        #######TYPE COUNT İSTİYORUM HER BİR POLEN TİPİNİN TOPLAM RESİM SAYISI
    lists = sorted(pollens_imgs.items()) # sorted by key, return a list of tuples
    x, y = zip(*lists) # unpack a list of pairs into two tuples
    print(x,y)
    plt.plot(x, y)
    plt.show()

    for key in classes:
        classes[key]["tot_img"] = 0
        for _type in classes[key]["types"]:
            classes[key]["tot_img"] += pollens_imgs[_type]


    print("##########################################")
    print("classes:", classes)

    tot_imgs = {}
    for key in classes:
        tot_imgs[key] = classes[key]["tot_img"]

    print("##########################################")
    print("tot_imgs:", tot_imgs)

    print("############################")
    print("pollens_imgs:", pollens_imgs)
    
    treemap(classes, pollens_imgs, tot_imgs)
    #treemap2()
